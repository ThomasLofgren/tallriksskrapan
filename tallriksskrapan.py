#!/usr/bin/python
# -*- coding: utf-8 -*-
import requests
import urllib.request
import io
import json
import logging
import cgi, cgitb 
import datetime

from lxml import html,etree
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine
from docx import Document

week_number = week_number = datetime.datetime.utcnow().isocalendar()[1]
lastWeek = "0"
daySearch = "fredag"

def getNextDay(day) :
    daysInWeek = [u'måndag', 'tisdag', 'onsdag', 'torsdag', 'fredag']
    retval = day
    getNext = False
    for weekDay in daysInWeek :
        if getNext :
            retval = weekDay
            break
        if weekDay == day :
            getNext = True
    return retval

def utf8text(text):
        return text.encode('raw_unicode_escape').decode('utf-8')

def parse_vecka():
    # week_number = datetime.datetime.utcnow().isocalendar()[1]
    return 'Det är nu vecka %s' % week_number

def parse_kompassen():
    ret = "### KOMPASSEN ###" + "\n"
    answer = requests.get('http://www.restaurangkompassen.se/index.php?option=com_content&view=article&id=64&Itemid=66')
    root = html.fromstring(answer.text)
    friday_found = False
    nextDay = getNextDay(daySearch)
    for child in root.xpath('//div[@class="screen"]/div/div/div'):
        if friday_found and child.text:
            if nextDay != daySearch and nextDay in child.text.lower() :
               break
            ret += child.text
        elif child.text and daySearch in child.text.lower():
            friday_found = True
    return ret

def parse_teknikparken():
    ret = "### TEKNIKPARKEN ###" + "\n"
    answer = requests.get('http://www.restaurangteknikparken.se/index.php?option=com_content&view=article&id=46')
    root = html.fromstring(answer.text)
    friday_found = False
    nextDay = getNextDay(daySearch)
    for child in root.xpath('//div[@class="screen"]/div/div/div'):
        if friday_found and child.text:
            if nextDay != daySearch and nextDay in child.text.lower() :
                break
            ret += child.text
        elif child.text and daySearch in child.text.lower():
            friday_found = True
    return ret

def parse_gs():
    ret = "### Gourmetservice ###" + "\n"
    answer = requests.get('http://www.geflegourmetservice.se/lunch.php')
    root = html.fromstring(answer.text)

    for child in root.xpath('//div[@class="left_holder"]/p')[1:3]:
        ret += child.text_content()
    return ret

def parse_sop():
    ret = "### Resturang Salt & Peppar ###" + "\n"
    answer = requests.get('http://www.restaurangsaltopeppar.com/Lunchmeny.html')
    root = html.fromstring(answer.text)
    parent = root
    DayFound = False
    BrCC = 0
    for child in root.xpath('//div[@align="center"]'):
        for span in child :
            if not DayFound :
                if span.tag == "span" and span.get("class") == "text-class-10" and daySearch in span.text.lower() :
                    DayFound = True
            else :
                if BrCC < 2 and span.tag == "span" :
                    if span.get("class") == "text-class-10" :
                        BrCC = 3
                    else :
                        BrCC = 0
                        if span.text :
                            ret += span.text + '\n'
                elif BrCC < 2 and span.tag == "br" :
                    BrCC += 1

                if BrCC >= 2 :
                   break;

    return ret

def parse_hemlingby():
    ret = "### HEMLINGBY ###" + "\n"
    answer = requests.get('http://www.gavle.se/Uppleva--gora/Idrott-motion-och-friluftsliv/Friluftsliv-och-motion/Hemlingby-friluftsomrade/Hemlingbystugan/Fika-och-ata/')
    root = html.fromstring(answer.text)
    for child in root.xpath('//a'):
        if child.text and "meny vecka" in child.text.lower() and str(week_number) in child.text.lower():
            hemlingby_link='http://www.gavle.se' + child.get('href')
            break
    textAsArray = parse_pdf(hemlingby_link)
    ret += getFoodFromPDFArray(textAsArray)
    return ret
    
#Takes url to pdf file and returns text split on newline into array
def parse_pdf(pdf_url):

    logging.propagate = False 
    logging.getLogger().setLevel(logging.ERROR)
    remote_file = urllib.request.urlopen(pdf_url).read()
    memory_file = io.BytesIO(remote_file)
    parser = PDFParser(memory_file)
    doc = PDFDocument()
    parser.set_document(doc)
    #Warning sometimes, error in pdf?
    doc.set_parser(parser)
    doc.initialize('')
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    ret = []
    # Process each page contained in the document.
    for pageIdx, page in enumerate(doc.get_pages()):
        ret.append([])
        interpreter.process_page(page)
        layout = device.get_result()
        for idx, lt_obj in enumerate(layout):
            if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                if len(lt_obj.get_text().strip()) > 0:
                    ret[pageIdx].append((lt_obj.get_text().splitlines()))
    return ret


def getFoodFromPDFArray(pdfArray):
    correctWeek = False
    ret = ""
    for page in pdfArray:
        for idx, line in enumerate(page):
            #Check if correct week
            if ("vecka " + str(week_number) + ":") in line[0]:
                correctWeek = True;
                continue;
            #If correct week and day is fredag
            elif correctWeek and daySearch in line[0].lower():

                #If the line is bigger than 1 it contains the food after 'fredag'
                if len(line) > 1:
                    for x in range(1, len(line)):
                        ret += line[x] + "\n"
                #The line only contained 'fredag' so the food is in the line after 'fredag'
                else:
                    line = page[idx+1]
                    for x in range(0, len(line)):
                        ret += line[x] + "\n"
                return ret

    return "Oops something went wrong"


def parse_gustafsbro():
    ret = "### Gustafsbro ###" + "\n"
    answer = requests.get('http://www.gavlelunch.se/gustafsbro.asp')
    root = html.fromstring(answer.text)
    friday_found = False

    #Get friday from table
    for weekdayTable in root.xpath('//body/font/table/tr[1]/td[1]/div/table'):
        for day in weekdayTable.xpath('tr[1]/td[1]/font/strong'):
            if day.text and daySearch in day.text.lower():
                friday_found = True
                break

    #If friday is found print food
    if friday_found:
           for food in weekdayTable.xpath('tr[2]/td[1]/font/ul/li'):
               ret += food.text.strip() + "\n"
    else:
           ret += "Oops something went wrong"
    return ret


def parse_sodersKalla():
    ret = "### Söders källa ###" + "\n"
    url = ""
    answer = requests.get('http://www.soderskalla.se/restaurangen/')
    root = html.fromstring(answer.text)

    #Get url for menu
    for child in root.xpath('//a'):
        if child.text and ("lunchmeny v" + str(week_number)) in child.text.lower():
            url = child.get('href')
            break
        #Check if menu has been updated from the week before
        elif child.text and ("lunchmeny v" + lastWeek) in child.text.lower():
             ret += "Menyn har ännu inte blivit uppdaterad"

    #Fetch document
    if url : 
        if not url.startswith('http:') :
            url = 'http:' + url
        answer = requests.get(url)
        memory_file = io.BytesIO(answer.content)
        doc = Document(memory_file)

        food = ""
        #Parse document and look for fredag, food is in the index after fredag
        for idx, para in enumerate(doc.paragraphs):
            if daySearch in para.text.lower():
                food = doc.paragraphs[idx+1].text + "\n"

        if food:
            ret += food
        else:
            ret += "Oops something went wrong"
    return ret

def parse_koket():
    ret = "### Köket ###" + "\n"

    answer = requests.get('http://koketlunch.se/meny.html')
    root = html.fromstring(answer.text)
    friday_found = False
    food = ""

    for line in root.xpath('//p/span'):
       #Get friday from table
        if line.text and daySearch in line.text.lower():
                friday_found = True

        if friday_found:
            if line.text.strip():
                #Fix encodings and remove '-' in the beginning of the different foods
                #Removed "fredag" printout to look more like the other printouts //Robert
                if daySearch in utf8text(line.text).lower():
                    pass
                    #ret += utf8text(line.text) + "\n"
                elif "stängt" in utf8text(line.text).lower():
                    food += utf8text(line.text) + "\n"
                else:
                    food += utf8text(line.text)[1:] + "\n"
            else:
                break
    if food:
        ret += food
    else:
        ret += "Oops something went wrong"
    return ret

def parse_kryddan():
    ret = "### Kryddan ###" + "\n"
    answer = requests.get('http://www.kryddan35.se/hem/')
    root = html.fromstring(answer.text)
    friday_found = False
    food = ""
    nextDay = getNextDay(daySearch)
    for child in root.xpath('//div[@id="veckans"]'):
        lines = child.text_content().split("\n")
        for line in lines:
            if friday_found and child.text:
                if nextDay != daySearch and nextDay in line.lower() :
                    break
                food += line + "\n"
            elif line and daySearch in line.lower():
                friday_found = True
    if food:
        ret += food
    else:
        ret += "Oops something went wrong"
    return ret

def get_jsonMenu(resturant):
    
    vecka = parse_vecka()
    restList = []
    if resturant == "teknikparken" :
        teknikparken = parse_teknikparken()
        restList.append({'restuarang':resturant, 'meny':teknikparken})
    elif resturant == "kompassen" :
        kompassen= parse_kompassen()
        restList.append({'restuarang':resturant,'meny':kompassen})
    elif resturant == "hemlingby" :
        hemlingby = parse_hemlingby()
        restList.append({'restuarang':resturant, 'meny':hemlingby})
    elif resturant == "gs" :
        gs = parse_gs()
        restList.append({'restuarang':resturant, 'meny':gs})
    elif resturant == "gustafsbro" :
        gustafsbro= parse_gustafsbro()
        restList.append({'restuarang':resturant, 'meny':gustafsbro})
    elif resturant == "soderskalla" :
        sodersKalla = parse_sodersKalla()
        restList.append({'restuarang':resturant, 'meny':soderskalla})
    elif resturant == "koket" :
        koket = parse_koket()
        restList.append({'restuarang':resturant, 'meny':koket})
    elif resturant == "kryddan" :
        kryddan = parse_kryddan()
        restList.append({'restuarang':resturant, 'meny':kryddan})
    elif resturant == "sop" :
        sop = parse_sop()
        restList.append({'restuarang':resturant, 'meny':sop})
    else :
        teknikparken = parse_teknikparken()
        kompassen= parse_kompassen()
        hemlingby = parse_hemlingby()
        gs = parse_gs()
        gustafsbro= parse_gustafsbro()
        sodersKalla = parse_sodersKalla()
        koket = parse_koket()
        kryddan = parse_kryddan()
        sop = parse_sop()
        restList = [{'restuarang': 'teknikparken', 'meny':teknikparken}, {'restuarang':'kompassen', 'meny':kompassen}, {'restuarang':'hemlingby', 'meny':hemlingby}, {'restuarang':'gs', 'meny':gs}, {'restuarang':'gustafsbro', 'meny':gustafsbro}, {'restuarang':'sodersKalla', 'meny':sodersKalla}, {'restuarang':'koket', 'meny':koket}, {'restuarang':'kryddan', 'meny':kryddan}, {'restuarang':'sop', 'meny':sop}]

    json_string = {'vecka':str(week_number), 'dag':daySearch, 'restuaranger':restList}
    return json.dumps(json_string)

def get_jsonResturants():
    json_resturants = {'restuaranger':[{'restuarang': 'teknikparken', 'text':'Teknikparken'}, {'restuarang':'kompassen', 'text': 'Kompassen'}, {'restuarang':'hemlingby', 'text':'Hemlingby'}, {'restuarang':'gs', 'text':'Gourmetservice'}, {'restuarang':'gustafsbro', 'text':'Gustafsbro'}, {'restuarang':'sodersKalla', 'text':'Söders Källa'}, {'restuarang':'koket', 'text':'Köket'}, {'restuarang':'kryddan', 'text':'Kryddan'}, {'restuarang':'sop', 'text':'Salt & Peppar'}]}
    return json.dumps(json_resturants)

def get_commandHelp() :
    json_help = {'commands':[{'name':'restuaranger', 'description':'Returns resturants in Swedish'},{'name':'menu', 'description':'Returns menus for all resturants, if parameter <resturant> with value of any resturant it will return that resturants menu. If paramater <dag> is set to any day [måndag,tisdag,onsdag,torsdag,fredag] it will return that day.'}]}
    return json.dumps(json_help)

def main():
    print("Content-Type: application/json\n\n")
    form = cgi.FieldStorage()
    command = form.getvalue("command")
    if command == "restuaranger" :
        print(get_jsonResturants())
    elif command == "menu" : 
        rest = form.getvalue("resturant")
        if form.getvalue("dag") :
            global daySearch
            daySearch = form.getvalue("dag")
        print(get_jsonMenu(rest))
    else :
        print(get_commandHelp())

    return
    
if __name__ == '__main__':
    main()
